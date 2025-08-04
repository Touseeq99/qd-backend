import os
import asyncio
import logging
from concurrent.futures import ThreadPoolExecutor
from typing import Dict, Any, List , Optional
from pydantic import SecretStr

# Configure logger
logger = logging.getLogger(__name__)
from qdrant_client.http import models
from langchain_community.document_loaders import DirectoryLoader
from langchain.embeddings.base import Embeddings
from sentence_transformers import SentenceTransformer

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_qdrant import QdrantVectorStore
from qdrant_client import QdrantClient, models
from qdrant_client.models import Distance, VectorParams, SparseVectorParams
from langchain_qdrant import QdrantVectorStore, FastEmbedSparse, RetrievalMode
from langchain.schema import Document
import re
from langchain_openai import OpenAIEmbeddings

# Thread pool executor for async ingestion
executor = ThreadPoolExecutor(max_workers=10)

# -------------------------
# ✅ Clause-based Chunking Strategy
# -------------------------

def extract_docx_text(file_path: str) -> List[str]:
    """Extract text from .docx file including paragraphs and tables."""
    from docx import Document
    doc = Document(file_path)
    content = []
    
    # Extract paragraphs
    content.extend(para.text.strip() for para in doc.paragraphs if para.text.strip())
    
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            # Join all cell texts in the row with tabs
            row_text = ' | '.join(cell.text.strip() for cell in row.cells)
            if row_text:
                content.append(row_text)
    
    return content

def is_clause_line(text: str) -> Optional[str]:
    """
    Match only lines that truly start with a clause pattern like '6.1', '6.2.3'
    and extract the top-level clause number like '6', '7', etc.
    """
    match = re.match(r"^(\d{1,2})(\.\d+)+\s", text)  # must be at least like '6.1 '
    return match.group(1) if match else None

def is_heading_line(text: str) -> bool:
    """
    Basic heuristic for detecting non-numbered headings.
    Accepts short, capitalized lines with no clause numbering.
    """
    return (
        not is_clause_line(text)
        and len(text.split()) <= 6
        and text[0].isupper()
        and text == text.title()
    )

def chunk_by_clause_or_heading(paragraphs: List[str]) -> List[Dict]:
    """
    Hybrid chunker:
    - If numbered clauses (1., 2.1, etc.) exist: chunks by top-level number
    - If no clauses: falls back to heading-based chunking
    """
    chunks = []
    current_chunk = []
    current_title = None
    clause_found = False

    for para in paragraphs:
        clause = is_clause_line(para)
        heading = is_heading_line(para)

        if clause:
            clause_found = True
            if clause != current_title:
                if current_chunk:
                    chunks.append({
                        "title": current_title or "Preamble",
                        "content": "\n".join(current_chunk).strip()
                    })
                    current_chunk = []
                current_title = clause
        elif heading and not clause_found:
            # Use headings as chunks only if no clauses are detected at all
            if current_chunk:
                chunks.append({
                    "title": current_title or "Section",
                    "content": "\n".join(current_chunk).strip()
                })
                current_chunk = []
            current_title = para

        current_chunk.append(para)

    if current_chunk:
        chunks.append({
            "title": current_title or "Final",
            "content": "\n".join(current_chunk).strip()
        })

    return chunks

def process_document(text: str, policy_name: str) -> List[Document]:
    """Process document using clause-based chunking for structured documents."""
    text = text.strip()
    if not text:
        return []

    from config import get_config
    config = get_config()
    
    # Split into paragraphs while preserving table rows
    paragraphs = []
    current_para = []
    
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        # Check if this is a table row (contains | or is a row of a table)
        is_table_row = '|' in line or (current_para and '|' in current_para[0])
        
        if is_table_row:
            # Add to current paragraph if it's empty or also a table row
            if not current_para or '|' in current_para[0]:
                current_para.append(line)
            else:
                # Start a new paragraph for the table
                if current_para:
                    paragraphs.append('\n'.join(current_para))
                current_para = [line]
        else:
            # Regular paragraph text
            if current_para and '|' in current_para[0]:
                # Flush the current table
                if current_para:
                    paragraphs.append('\n'.join(current_para))
                current_para = [line]
            else:
                current_para.append(line)
    
    # Add the last paragraph
    if current_para:
        paragraphs.append('\n'.join(current_para))
    
    # Use enhanced chunking that handles both numbered clauses and headings
    chunks = chunk_by_clause_or_heading(paragraphs)
    
    # Convert to Document objects with metadata
    final_chunks = []
    for i, chunk in enumerate(chunks):
        final_chunks.append(Document(
            page_content=chunk["content"],
            metadata={
                "heading": chunk["title"],
                "chunk_type": "section",
                "policy": policy_name,
                "chunking_method": "heading_with_clauses",
                "chunk_index": i,
                "content_length": len(chunk["content"])
            }
        ))
    
    return final_chunks


async def delete_and_recreate_collection(
    client: Any,
    collection_name: str,
    vector_size: int = 768
) -> None:
    """Delete and recreate a collection with the specified configuration."""
    
    try:
        client.delete_collection(collection_name=collection_name)
        logger.info(f"Deleted existing collection: {collection_name}")
    except Exception as e:
        if "doesn't exist" not in str(e):
            logger.warning(f"Error deleting collection: {e}")
    
    # Create new collection with named vectors
    logger.info(f"Creating new collection: {collection_name}")
    client.create_collection(
        collection_name=collection_name,
        vectors_config={
            "dense": models.VectorParams(
                size=vector_size,
                distance=models.Distance.COSINE
            )
        },
        sparse_vectors_config={
            "sparse": models.SparseVectorParams(
                index=models.SparseIndexParams(on_disk=False)
            )
        },
        hnsw_config=models.HnswConfigDiff(ef_construct=64)
    )
    logger.info(f"Created new collection with named vectors: dense and sparse")


async def ingest_documents_to_qdrant_async(
    directory_path: str,
    qdrant_url: str,
    qdrant_api: str,
    collection_name: str = "HRDOCS",
    chunk_size: int = 1000,
    chunk_overlap: int = 200
) -> Dict[str, Any]:
    """
    Asynchronously ingests documents from a folder into Qdrant.

    Args:
        directory_path (str): Local folder path containing PDF/DOCX files.
        qdrant_url (str): Qdrant server URL.
        collection_name (str): Name of the Qdrant collection.
        qdrant_api (str): Qdrant API key (if any).

    Returns:
        dict: Ingestion status.
    """

    if not os.path.exists(directory_path):
        return {"status": "error", "message": "Directory path does not exist."}

    def sync_ingest():
        # Import config here to avoid circular imports
        from config import get_config
        config = get_config()
        logger = logging.getLogger(__name__)
        
        loader = DirectoryLoader(
            path=directory_path,
            glob="**/*",
        )
        documents = loader.load()

        logger.info(f"DirectoryLoader found {len(documents)} documents in {directory_path}")
        
        if not documents:
            logger.error(f"No documents found in directory: {directory_path}")
            return {"status": "No documents found"}
        
        # Log document details for debugging
        for i, doc in enumerate(documents):
            logger.info(f"Document {i+1}: {doc.metadata.get('source', 'unknown')}")
            logger.info(f"  Content length: {len(doc.page_content)}")
            logger.info(f"  Content preview: {doc.page_content[:100]}...")

        # Use clause-based chunking strategy for better retrieval
        all_chunks = []
        logger.info(f"Processing {len(documents)} documents for chunking...")
        
        for i, doc in enumerate(documents):
            policy_name = doc.metadata.get('source', 'unknown')
            # Extract filename from path for better policy identification
            if 'source' in doc.metadata:
                policy_name = os.path.basename(doc.metadata['source'])
            
            logger.info(f"Processing document {i+1}/{len(documents)}: {policy_name}")
            logger.info(f"Document content length: {len(doc.page_content)} characters")
            logger.info(f"First 200 characters: {doc.page_content[:200]}...")
            
            # Process each document with clause-based chunking
            doc_chunks = process_document(doc.page_content, policy_name)
            logger.info(f"Generated {len(doc_chunks)} chunks for document: {policy_name}")
            
            # Filter out empty chunks
            valid_chunks = [chunk for chunk in doc_chunks if chunk.page_content.strip()]
            logger.info(f"Valid chunks (non-empty): {len(valid_chunks)}")
            
            all_chunks.extend(valid_chunks)
        
        chunks = all_chunks
        logger.info(f"Total chunks generated: {len(chunks)}")
        
        if not chunks:
            logger.warning("No valid chunks generated! Falling back to simple text splitting...")
            # Fallback to simple text splitting
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=800,
                chunk_overlap=100,
                separators=["\n\n", "\n", ".", " ", ""]
            )
            chunks = splitter.split_documents(documents)
            logger.info(f"Fallback chunks generated: {len(chunks)}")

        # Initialize SentenceTransformer for embeddings
        logger.info("Initializing SentenceTransformer with BAAI/bge-small-en-v1.5 model")
        
        class SentenceTransformerEmbeddings(Embeddings):
            def __init__(self, model_name: str = "BAAI/bge-base-en-v1.5"):
                self.model = SentenceTransformer(model_name)
                
            def embed_documents(self, texts: List[str]) -> List[List[float]]:
                """Embed a list of documents using SentenceTransformer."""
                if not texts:
                    return []
                    
                # Convert input to list of strings if it's a single string
                if isinstance(texts, str):
                    texts = [texts]
                    
                try:
                    # Generate embeddings in batches to handle large inputs
                    batch_size = 32
                    all_embeddings = []
                    
                    for i in range(0, len(texts), batch_size):
                        batch = texts[i:i + batch_size]
                        embeddings = self.model.encode(
                            batch,
                            convert_to_numpy=True,
                            normalize_embeddings=True
                        )
                        all_embeddings.extend(embeddings.tolist())
                    
                    return all_embeddings
                    
                except Exception as e:
                    logger.error(f"Error embedding documents: {str(e)}")
                    raise
                    
            def embed_query(self, text: str) -> List[float]:
                """Embed a query using SentenceTransformer with the required prefix."""
                try:
                    # Format the query with the required prefix
                    formatted_query = "Represent this sentence for searching relevant passages: " + text
                    embedding = self.model.encode(
                        formatted_query,
                        convert_to_numpy=True,
                        normalize_embeddings=True
                    )
                    return embedding.tolist()
                    raise ValueError('No valid embedding returned')
                except Exception as e:
                    logger.error(f"Error embedding query: {str(e)}")
                    raise
        
        # Initialize SentenceTransformer embeddings
        llm_config = config.get_active_llm_config() 
        if llm_config['provider']=="openai":
            embeddings = OpenAIEmbeddings(model_name="text-embedding-3-large",dimensions=768)
        else:
            embeddings = SentenceTransformerEmbeddings(
                model_name=config.embedding.sentence_transformer_model
            )   


        client = QdrantClient(url=qdrant_url, api_key=qdrant_api, timeout=120)

        # Initialize sparse embeddings for hybrid search
        sparse_embeddings = FastEmbedSparse(model_name="Qdrant/bm25")
        
        # Always use named vectors for consistency
        vector_name = "dense"
        sparse_vector_name = "sparse"
        
        # Initialize vector store with explicit configuration
        logger.info(f"Initializing QdrantVectorStore with vector_name='{vector_name}', sparse_vector_name='{sparse_vector_name}'")
        
        vectordatabase = QdrantVectorStore(
            client=client,
            collection_name=collection_name,
            embedding=embeddings,
            sparse_embedding=sparse_embeddings if sparse_vector_name else None,
            vector_name=vector_name,
            sparse_vector_name=sparse_vector_name
        )
        
        # Add documents to the vector store
        logger.info(f"Adding {len(chunks)} chunks to the vector store...")
        vectordatabase.add_documents(chunks)
        logger.info("Successfully added documents to the vector store")
  
        return {"status": "success", "chunks_ingested": len(chunks)}

    return await asyncio.get_event_loop().run_in_executor(executor, sync_ingest)


def check_or_create_qdrant_collection(
    qdrant_url: str, 
    qdrant_api: str,
    collection_name: str,
    vector_size: int = 768,
    distance_metric = None,  # Will be set to models.Distance.COSINE if None
    sparse_vectors: bool = True
) -> str:
    """
    Creates a Qdrant collection with the specified configuration.
    
    Args:
        qdrant_url: URL of the Qdrant server
        qdrant_api: API key for Qdrant
        collection_name: Name of the collection to create
        vector_size: Size of the dense vectors (default: 768)
        distance_metric: Distance metric for vector search (default: COSINE)
        sparse_vectors: Whether to enable sparse vectors (default: True)
        
    Returns:
        str: Status message indicating the action taken
    """
    from qdrant_client import QdrantClient, models
    
    # Set default distance metric if not provided
    if distance_metric is None:
        distance_metric = models.Distance.COSINE
    
    try:
        # Initialize client with timeout
        client = QdrantClient(
            url=qdrant_url, 
            api_key=qdrant_api,
            timeout=30.0  # 30 seconds timeout
        )
        
        # Delete collection if it exists
        try:
            client.delete_collection(collection_name=collection_name)
            logger.info(f"Deleted existing collection: {collection_name}")
        except Exception as e:
            if "doesn't exist" not in str(e):
                logger.warning(f"Error deleting collection (may not exist): {e}")
        
        # Prepare vector configuration
        vectors_config = {
            "dense": models.VectorParams(
                size=vector_size,
                distance=distance_metric
            )
        }
        
        # Prepare sparse vectors config if enabled
        sparse_vectors_config = {
            "sparse": models.SparseVectorParams()
        } if sparse_vectors else None
        
        # Create the collection
        client.create_collection(
            collection_name=collection_name,
            vectors_config=vectors_config,
            sparse_vectors_config=sparse_vectors_config,
            hnsw_config=models.HnswConfigDiff(ef_construct=64)
        )
        
        logger.info(f"Successfully created collection: {collection_name}")
        return f"Created new collection '{collection_name}' with {'dense' + (' and sparse' if sparse_vectors else '')} vectors"
            
    except Exception as e:
        error_msg = f"Error creating collection: {str(e)}"
        logger.error(error_msg)
        raise Exception(error_msg)
    # Import config here to avoid circular imports
    from config import get_config
    config = get_config() 
    
    # Use config vector size if default is provided
    if vector_size == 768:  # Default value
        vector_size = config.database.vector_size
    
    client = QdrantClient(url=qdrant_url, api_key=qdrant_api)
    existing = [c.name for c in client.get_collections().collections]

    if collection_name in existing:
        print(f"✅ Collection '{collection_name}' already exists.")
    else:
        client.create_collection(
            collection_name=collection_name,
            vectors_config=VectorParams(size=vector_size, distance=distance_metric),
        )
        print(f"✅ Collection '{collection_name}' created.")
    return collection_name


def delete_qdrant_collection(qdrant_url: str, qdrant_api: str, collection_name: str) -> None:
    """
    Deletes a Qdrant collection if it exists.
    """
    client = QdrantClient(url=qdrant_url, api_key=qdrant_api)
    existing = [c.name for c in client.get_collections().collections]

    if collection_name in existing:
        client.delete_collection(collection_name=collection_name)
        print(f"🗑️  Collection '{collection_name}' has been deleted.")
    else:
        print(f"⚠️  Collection '{collection_name}' does not exist. Nothing to delete.")
        
from qdrant_client import QdrantClient

def get_collections_with_chunk_counts(qdrant_url: str, qdrant_api: str) -> dict:
    """
    Returns all Qdrant collections along with their vector (chunk) count.

    Output:
    {
        "collections": [
            {"name": "my_collection", "chunks": 132},
            {"name": "resumes_batch_01", "chunks": 58},
            ...
        ]
    }
    """
    client = QdrantClient(url=qdrant_url, api_key=qdrant_api)
    collections = client.get_collections().collections

    results = []
    for collection in collections:
        name = collection.name
        try:
            count = client.get_collection(collection_name=name).points_count
            
        except Exception:
            count = 0
        results.append({"name": name, "chunks": count})

    return {"collections": results}

